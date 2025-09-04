#!/usr/bin/env python
from __future__ import annotations

import re
import random
import argparse
from typing import TYPE_CHECKING, List, Dict, Optional, Union, Tuple, Literal, TypedDict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if TYPE_CHECKING:  # Imported only for type checking; no runtime dependency changes
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph


class MatchingQuestion(TypedDict):
    left: str
    right: str


class TextQuestion(TypedDict, total=False):
    text: str
    lines: int  # optional


class BlankQuestion(TypedDict):
    text: str


class OralQuestion(TypedDict):
    text: str
    sub_points: List[str]


SectionType = Literal["short", "long", "matching_v", "matching_h", "blanks", "oral"]


class Section(TypedDict):
    name: str
    questions: List[Union[MatchingQuestion, TextQuestion, BlankQuestion, OralQuestion]]
    type: Optional[SectionType]
    separate_sheet: bool
    subtitle: Optional[str]


class TestData(TypedDict):
    subject: str
    title: str
    sections: List[Section]


def parse_template(file_path: str) -> TestData:
    def _parse_section(
        line: str, current_section: Optional[Section], sections: List[Section]
    ) -> Section:
        name = line.split(":", 1)[1].strip()
        if current_section:
            sections.append(current_section)
        return {
            "name": name,
            "questions": [],
            "type": None,
            "separate_sheet": False,
            "subtitle": None,
        }

    def _parse_question(
        q: str, current_section: Section
    ) -> Optional[Union[MatchingQuestion, TextQuestion, BlankQuestion, OralQuestion]]:
        if current_section["type"] in ["matching_v", "matching_h"]:
            if "->" in q:
                left, right = [x.strip() for x in q.split("->")]
                return {"left": left, "right": right}
        elif current_section["type"] == "blanks":
            return {"text": q}
        elif current_section["type"] == "oral":
            # For oral questions, we'll handle sub-points separately
            return {"text": q, "sub_points": []}
        else:
            lines_count: Optional[int] = None
            # Remove '(N line[s])' from the question text if present
            match = re.search(r"\((\d+)\s+lines?\)", q)
            if match:
                lines_count = int(match.group(1))
                text = re.sub(r"\(\d+\s+lines?\)", "", q).strip()
            else:
                text = q.strip()
            # Only include 'lines' key if present to match TextQuestion total=False
            if lines_count is not None:
                return {"text": text, "lines": lines_count}
            return {"text": text}

    with open(file_path, "r") as f:
        lines: List[str] = f.readlines()

    test: TestData = {"subject": "", "title": "", "sections": []}
    sections: List[Section] = []
    current_section: Optional[Section] = None

    for line in lines:
        original_line = line
        line = line.rstrip()  # Only strip trailing whitespace, preserve indentation
        if not line.strip():  # Check if line is empty after stripping all whitespace
            continue
        if line.startswith("# Test"):
            continue
        elif line.startswith("Subject:"):
            test["subject"] = line.split(":", 1)[1].strip()
        elif line.startswith("Title:"):
            test["title"] = line.split(":", 1)[1].strip()
        elif line.startswith("## Section:"):
            current_section = _parse_section(line, current_section, sections)
        elif line.startswith("Type:"):
            if current_section:
                current_section["type"] = line.split(":", 1)[1].strip()  # type: ignore[assignment]
        elif line.startswith("Subtitle:"):
            if current_section:
                current_section["subtitle"] = line.split(":", 1)[1].strip()
        elif line.startswith("Oral:"):
            # Legacy support - convert "Oral: yes" to Type: oral for backward compatibility
            if current_section:
                is_oral = line.split(":", 1)[1].strip().lower() == "yes"
                if is_oral:
                    current_section["type"] = "oral"  # type: ignore[assignment]
        elif line.startswith("Separate Sheet:"):
            if current_section:
                current_section["separate_sheet"] = (
                    line.split(":", 1)[1].strip().lower() == "yes"
                )
        elif line.startswith("    - ") or line.startswith("\t- "):
            # Handle sub-points for oral questions (indented with 4 spaces or tab)
            if (
                current_section
                and current_section["type"] == "oral"
                and current_section["questions"]
            ):
                sub_point = line.strip()[2:].strip()  # Remove "- " from the beginning
                last_question = current_section["questions"][-1]
                if isinstance(last_question, dict) and "sub_points" in last_question:
                    last_question["sub_points"].append(sub_point)  # type: ignore[index]
        elif line.startswith("- "):
            if current_section:
                q = line[2:].strip()
                question = _parse_question(q, current_section)
                if question:
                    current_section["questions"].append(question)

    if current_section:
        sections.append(current_section)

    test["sections"] = sections
    return test


def generate_docx(
    test_data: TestData,
    output_path: str = "test.docx",
    template_path: Optional[str] = None,
) -> None:
    def _remove_leading_empty_paragraph(doc: DocxDocument) -> None:
        # python-docx starts a new document with a single empty body paragraph.
        # If present (and empty/whitespace-only), remove it to avoid a blank line
        # before the first real content we add below.
        try:
            if doc.paragraphs and not doc.paragraphs[0].text.strip():
                p = doc.paragraphs[0]._element
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)
        except Exception:
            # Non-fatal: if removal fails, continue; later content will still render.
            pass

    def _add_markdown_run(paragraph: Paragraph, text: str) -> None:
        # Simple Markdown parser for **bold**, *italic*, _italic_
        pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)"
        pos = 0
        for match in re.finditer(pattern, text):
            start, end = match.span()
            if start > pos:
                paragraph.add_run(text[pos:start])
            md = text[start:end]
            run = paragraph.add_run(md.strip("*_"))
            if md.startswith("**") and md.endswith("**"):
                run.bold = True
            elif (md.startswith("*") and md.endswith("*")) or (
                md.startswith("_") and md.endswith("_")
            ):
                run.italic = True
            pos = end
        if pos < len(text):
            paragraph.add_run(text[pos:])

    def _set_page_layout(doc: DocxDocument, section=None) -> None:
        if section is None:
            section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = section.top_margin = (
            section.bottom_margin
        ) = Inches(0.75)
        section.different_first_page_header_footer = True

    def _add_header(doc: DocxDocument, section=None) -> None:
        if section is None:
            section = doc.sections[0]
        first_header = section.first_page_header
        header_p = first_header.paragraphs[0]
        header_p.style.paragraph_format.space_before = Pt(27)
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Automated en-space generation for Name and Date fields
        def add_label_and_blank(paragraph, label, num_spaces):
            run = paragraph.add_run("\u2000" + label + "\u2000")
            blank_run = paragraph.add_run("\u2000" * num_spaces)
            blank_run.underline = True

        add_label_and_blank(header_p, "Name:", 18)
        add_label_and_blank(header_p, "Date:", 12)

    def _add_subject_and_title(doc: DocxDocument, test_data: TestData) -> None:
        doc.add_paragraph(f"{test_data['subject']} Test", style="Subtitle")
        doc.add_paragraph(test_data["title"], style="Title")

    def _add_section(doc: DocxDocument, section: Section) -> None:
        doc.add_heading(section["name"], level=2)
        heading_p = doc.paragraphs[-1]
        if section["type"] == "long" and section["separate_sheet"]:
            heading_p.paragraph_format.space_after = Pt(0)
            p = doc.add_paragraph("Use a separate sheet of paper")
            p.paragraph_format.space_before = Pt(0)
            run = p.runs[0]
            run.italic = True
            run.font.size = Pt(10)
        if section["type"] == "oral":
            heading_p.paragraph_format.space_after = Pt(0)
            p = doc.add_paragraph("To be completed orally")
            p.paragraph_format.space_before = Pt(0)
            run = p.runs[0]
            run.italic = True
            run.font.size = Pt(10)
        # Add optional subtitle rendered similarly to the separate-sheet note
        subtitle = section.get("subtitle")
        if subtitle:
            heading_p.paragraph_format.space_after = Pt(0)
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(0)
            # Render markdown into the subtitle paragraph
            _add_markdown_run(p_sub, subtitle)
            # Apply the same styling as the separate-sheet note to all runs
            for run_sub in p_sub.runs:
                run_sub.italic = True
                run_sub.font.size = Pt(10)

    def _add_short_questions(doc: DocxDocument, section: Section) -> None:
        for q in section["questions"]:  # type: ignore[assignment]
            p = doc.add_paragraph(style="List Number")
            _add_markdown_run(p, q["text"])  # type: ignore[index]
            p.paragraph_format.line_spacing = 1.0  # Single-spaced question text
            num_lines = q["lines"] if isinstance(q, dict) and q.get("lines") else 1  # type: ignore[index]
            for _ in range(num_lines):
                blank_p = doc.add_paragraph()
                blank_p.add_run("\t")
                blank_p.paragraph_format.tab_stops.add_tab_stop(Inches(7))
                run = blank_p.runs[-1]
                run.underline = True
                blank_p.paragraph_format.line_spacing = 1.5

    def _add_long_questions(doc: DocxDocument, section: Section) -> None:
        for q in section["questions"]:  # type: ignore[assignment]
            p = doc.add_paragraph(style="List Number")
            _add_markdown_run(p, q["text"])  # type: ignore[index]
            p.paragraph_format.line_spacing = 1.0  # Single-spaced question text
            if not section["separate_sheet"]:
                num_lines = q.get("lines", 10)  # type: ignore[attribute-defined-outside-init]
                for _ in range(num_lines):
                    blank_p = doc.add_paragraph()
                    blank_p.add_run("\t")
                    blank_p.paragraph_format.tab_stops.add_tab_stop(Inches(7))
                    run = blank_p.runs[-1]
                    run.underline = True
                    blank_p.paragraph_format.line_spacing = 1.5

    def _add_matching_v(doc: DocxDocument, section: Section) -> None:
        pairs: List[Tuple[str, str]] = [
            (q["left"], q["right"]) for q in section["questions"]  # type: ignore[index]
        ]
        random.shuffle(pairs)
        rights: List[str]
        lefts: List[str]
        if pairs:
            lefts_t, rights_t = zip(*pairs)
            lefts, rights = list(lefts_t), list(rights_t)
        else:
            lefts, rights = [], []
        rights = list(rights)
        random.shuffle(rights)
        doc_section = doc.sections[0]
        usable_width = (
            doc_section.page_width - doc_section.left_margin - doc_section.right_margin
        )
        display_rights: List[str] = [
            f"{chr(65 + i)}. {r}" for i, r in enumerate(rights)
        ]
        max_right_len: int = max((len(s) for s in display_rights), default=0)
        char_pt: float = 6.5
        right_width = Pt(max_right_len * char_pt + 12)
        max_right_allowed = usable_width * 0.6
        if right_width > max_right_allowed:
            right_width = max_right_allowed
        table = doc.add_table(rows=len(lefts), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if len(table.columns) >= 2:
            table.columns[1].width = right_width
            left_width = usable_width - right_width
            if left_width <= Pt(0):
                left_width = usable_width - Pt(40)
            table.columns[0].width = left_width
        for i in range(len(lefts)):
            row_cells = table.rows[i].cells
            row_cells[0].text = ""
            left_p = row_cells[0].paragraphs[0]
            blank_run = left_p.add_run("\u2003\u2003")
            blank_run.underline = True
            left_p.add_run(f" {lefts[i]}")
            row_cells[1].text = ""
            right_p = row_cells[1].paragraphs[0]
            right_text = (
                display_rights[i]
                if i < len(display_rights)
                else f"{chr(65 + i)}. {rights[i]}"
            )
            right_p.add_run(right_text)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)

    def _add_matching_h(doc: DocxDocument, section: Section) -> None:
        pairs: List[Tuple[str, str]] = [
            (q["left"], q["right"]) for q in section["questions"]  # type: ignore[index]
        ]
        terms: List[str] = []
        defs: List[str] = []
        if pairs:
            terms_t, defs_t = zip(*pairs)
            terms = list(terms_t)
            defs = list(defs_t)
        n_terms = len(terms)
        if n_terms:
            best_rows = 1
            best_empty = None
            for r in range(1, 4):
                cols = (n_terms + r - 1) // r
                empty = r * cols - n_terms
                if (
                    best_empty is None
                    or empty < best_empty
                    or (empty == best_empty and r < best_rows)
                ):
                    best_rows = r
                    best_empty = empty
            rows = best_rows
            cols = (n_terms + rows - 1) // rows
            term_bank = doc.add_table(rows=rows, cols=cols)
            term_bank.alignment = WD_TABLE_ALIGNMENT.CENTER
            term_bank.autofit = True
            labels = [f"{chr(65 + i)}. {t}" for i, t in enumerate(terms)]
            idx = 0
            for r in range(rows):
                for c in range(cols):
                    cell = term_bank.rows[r].cells[c]
                    cell.text = ""
                    if idx < len(labels):
                        p = cell.paragraphs[0]
                        run = p.add_run(labels[idx])
                        run.font.size = Pt(10)
                        idx += 1
            for row in term_bank.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(0)
        n = len(defs)
        if n == 0:
            return
        rows = (n + 1) // 2
        match_table = doc.add_table(rows=rows, cols=4)
        match_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        match_table.autofit = False
        for i in range(n):
            row = i // 2
            col = (i % 2) * 2
            blank_cell = match_table.rows[row].cells[col]
            blank_cell.text = ""
            p_blank = blank_cell.paragraphs[0]
            p_blank.paragraph_format.space_before = Pt(5.5)
            blank_run = p_blank.add_run("\u2003\u2003\u2003\u2003\u2003")
            blank_run.underline = True
            p_blank.add_run(" ")
            def_cell = match_table.rows[row].cells[col + 1]
            def_cell.text = ""
            p_def = def_cell.paragraphs[0]
            p_def.paragraph_format.space_before = Pt(5.5)
            p_def.add_run(defs[i])
        font_size_pt = 11
        num_em_spaces = 5
        emu_per_pt = 12700
        padding_inch = 0.16
        emu_per_inch = 914400
        padding_emu = int(padding_inch * emu_per_inch)
        blank_w = int(font_size_pt * num_em_spaces * emu_per_pt) + padding_emu
        doc_section = doc.sections[0]
        usable_width = (
            doc_section.page_width - doc_section.left_margin - doc_section.right_margin
        )
        def_w = int((usable_width - 2 * blank_w) / 2)
        match_table.columns[0].width = blank_w
        match_table.columns[1].width = def_w
        match_table.columns[2].width = blank_w
        match_table.columns[3].width = def_w
        for row in match_table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)

    def _add_blanks_questions(doc: DocxDocument, section: Section) -> None:
        for q in section["questions"]:
            p = doc.add_paragraph(style="List Number")
            text = q["text"]  # type: ignore[index]
            p.paragraph_format.line_spacing = 2.0  # Double-spaced for blanks section
            parts = re.split(r"(_+)", text)
            for part in parts:
                if part and part[0] == "_":
                    for _ in part:
                        run = p.add_run("\u2003")
                        run.underline = True
                else:
                    _add_markdown_run(p, part)

    def _add_oral_questions(
        doc: DocxDocument, section: Section, test_data: TestData
    ) -> None:
        # Add questions to main document without blank lines (already handled in _add_section)
        for q in section["questions"]:  # type: ignore[assignment]
            p = doc.add_paragraph(style="List Number")
            _add_markdown_run(p, q["text"])  # type: ignore[index]
            p.paragraph_format.line_spacing = 1.0  # Single-spaced question text
            # Sub-points should NOT appear on the test page - only on assessment sheet

        # Generate oral assessment sheet on a separate page
        _add_oral_assessment_sheet(doc, section, test_data)

    def _set_cell_border(cell, **kwargs):
        """
        Set cell's border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "val": "single", "color": "#00FF00", "space": "0"},
            start={"sz": 24, "val": "dashed", "color": "#0000FF", "space": "0"},
            end={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        )
        """
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
            if edge in kwargs:
                tag = f"w:{edge}"
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                for k, v in kwargs[edge].items():
                    element.set(qn(f"w:{k}"), str(v))

    def _add_oral_assessment_sheet(
        doc: DocxDocument, section: Section, test_data: TestData
    ) -> None:
        # Add a page break to start the oral assessment sheet
        doc.add_page_break()

        # Add a new section for the assessment sheet
        doc.add_section()

        # Set page layout for the new section
        _set_page_layout(doc, doc.sections[-1])

        # Add header to the new section
        _add_header(doc, doc.sections[-1])

        # Add header for the oral assessment sheet
        h = doc.add_heading(f"{test_data['title']} - Oral Assessment Sheet", level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Calculate number of rows needed
        total_rows = 0
        for q in section["questions"]:  # type: ignore[assignment]
            total_rows += 1  # Main question
            if isinstance(q, dict) and "sub_points" in q:
                total_rows += len(q["sub_points"])  # type: ignore[index]
        total_rows += 1  # Total row

        # Create table with 2 columns
        table = doc.add_table(rows=total_rows, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Set column widths - first column wider for questions, second narrower for scoring
        doc_section = doc.sections[0]
        usable_width = (
            doc_section.page_width - doc_section.left_margin - doc_section.right_margin
        )
        # Calculate minimum width for score column (4 em-spaces + padding)
        font_size_pt = 12
        num_em_spaces = 4
        emu_per_pt = 12700
        padding_inch = 0.1
        emu_per_inch = 914400
        score_width_emu = int(font_size_pt * num_em_spaces * emu_per_pt) + int(
            padding_inch * emu_per_inch
        )
        score_width = Pt(score_width_emu / emu_per_pt)
        question_width = usable_width - score_width

        table.columns[0].width = question_width
        table.columns[1].width = score_width

        # Populate the table
        row_idx = 0
        for q in section["questions"]:  # type: ignore[assignment]
            # Main question row
            question_cell = table.rows[row_idx].cells[0]
            score_cell = table.rows[row_idx].cells[1]

            question_cell.text = ""
            question_p = question_cell.paragraphs[0]
            _add_markdown_run(question_p, q["text"])  # type: ignore[index]

            score_cell.text = ""
            score_p = score_cell.paragraphs[0]
            # Add underlined em-space for scoring only if no sub-points
            if not (isinstance(q, dict) and "sub_points" in q and q["sub_points"]):
                score_run = score_p.add_run("\u2003\u2003\u2003\u2003")
                score_run.underline = True

            row_idx += 1

            # Sub-point rows (if any)
            if isinstance(q, dict) and "sub_points" in q:
                for i, sub_point in enumerate(q["sub_points"]):  # type: ignore[index]
                    sub_question_cell = table.rows[row_idx].cells[0]
                    sub_score_cell = table.rows[row_idx].cells[1]

                    sub_question_cell.text = ""
                    sub_question_p = sub_question_cell.paragraphs[0]
                    # Add tab indentation for sub-points
                    sub_question_p.add_run("\t")
                    _add_markdown_run(sub_question_p, sub_point)

                    sub_score_cell.text = ""
                    sub_score_p = sub_score_cell.paragraphs[0]
                    # Add underlined em-space for scoring
                    sub_score_run = sub_score_p.add_run("\u2003\u2003\u2003\u2003")
                    sub_score_run.underline = True

                    if i > 0:
                        for cell in table.rows[row_idx].cells:
                            _set_cell_border(cell, top={"color": "D3D3D3"})

                    row_idx += 1

        # Add notes row instead of total scoring
        notes_row = table.rows[row_idx]
        # Merge both columns for notes section
        notes_cell = notes_row.cells[0].merge(notes_row.cells[1])
        # Add Notes heading
        notes_cell.text = ""
        p_notes = notes_cell.paragraphs[0]
        run_notes = p_notes.add_run("Notes")
        run_notes.bold = True
        run_notes.font.size = Pt(9)
        # Reserve space for notes (5 blank lines)
        for _ in range(5):
            notes_cell.add_paragraph()

        # Format table cells
        for row in table.rows:
            for cell in row.cells:
                # Set cell padding to 4pt on all sides
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.first_child_found_in("w:tcMar")
                if tcMar is None:
                    tcMar = OxmlElement("w:tcMar")
                    tcPr.append(tcMar)
                padding_dxa = int(Pt(4).emu // 635)
                for side in ["top", "bottom", "left", "right"]:
                    tag = f"w:{side}"
                    element = tcMar.find(qn(tag))
                    if element is None:
                        element = OxmlElement(tag)
                        tcMar.append(element)
                    element.set(qn("w:w"), str(padding_dxa))
                    element.set(qn("w:type"), "dxa")
                # Remove paragraph after-spacing
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)

    def _finalize_paragraphs(doc: DocxDocument) -> None:
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text and paragraph.text.strip().startswith(
                "Use a separate sheet"
            ):
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                if idx > 0:
                    prev = doc.paragraphs[idx - 1]
                    try:
                        prev.paragraph_format.space_after = Pt(0)
                    except Exception:
                        pass
                continue

    doc = Document(template_path) if template_path else Document()
    _set_page_layout(doc)
    _remove_leading_empty_paragraph(doc)
    _add_subject_and_title(doc, test_data)
    for section in test_data["sections"]:
        _add_section(doc, section)
        typ = section["type"]
        if typ == "short":
            _add_short_questions(doc, section)
        elif typ == "long":
            _add_long_questions(doc, section)
        elif typ == "matching_v":
            _add_matching_v(doc, section)
        elif typ == "matching_h":
            _add_matching_h(doc, section)
        elif typ == "blanks":
            _add_blanks_questions(doc, section)
        elif typ == "oral":
            _add_oral_questions(doc, section, test_data)
    _finalize_paragraphs(doc)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a DOCX test file from a template"
    )
    parser.add_argument("input", help="Path to the input template file (.txt)")
    parser.add_argument(
        "-o",
        "--output",
        default="test.docx",
        help="Path to the output DOCX file (default: test.docx)",
    )
    parser.add_argument(
        "-t",
        "--template",
        default=None,
        help="Path to a DOCX template file (optional)",
    )
    args: argparse.Namespace = parser.parse_args()

    template_data = parse_template(args.input)
    generate_docx(template_data, args.output, template_path=args.template)
    print(f"Generated DOCX file: {args.output}")


if __name__ == "__main__":
    main()
