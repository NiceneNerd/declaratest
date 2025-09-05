#!/usr/bin/env python3
"""
Simple verification script to check that both Python and Rust versions 
generate DOCX files with similar structure and content.
"""

import sys
from docx import Document

def analyze_docx(filename):
    """Analyze a DOCX file and return basic structure info."""
    try:
        doc = Document(filename)
        
        info = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'has_header': bool(doc.sections[0].header.paragraphs),
            'content_types': set(),
            'text_samples': []
        }
        
        # Collect content types and samples
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                info['text_samples'].append(text[:50])  # First 50 chars
                
                # Identify content types
                if 'Test' in text and ('General Knowledge' in text or 'Wonders' in text):
                    info['content_types'].add('title')
                elif text.startswith('Quick Facts') or text.startswith('Creative Writing'):
                    info['content_types'].add('section_heading')
                elif 'Use the word' in text or 'What is the most common' in text:
                    info['content_types'].add('question')
                elif 'Use a separate sheet' in text or 'To be completed orally' in text:
                    info['content_types'].add('instruction')
                elif 'Name:' in text and 'Date:' in text:
                    info['content_types'].add('header')
                elif 'Oral Assessment Sheet' in text:
                    info['content_types'].add('assessment_sheet')
        
        # Check tables for matching questions and oral assessments
        if info['tables'] > 0:
            info['content_types'].add('tables')
            
        return info
        
    except Exception as e:
        return {'error': str(e)}

def main():
    if len(sys.argv) != 3:
        print("Usage: verify_docx.py <python_docx> <rust_docx>")
        sys.exit(1)
    
    python_file = sys.argv[1]
    rust_file = sys.argv[2]
    
    print("Analyzing DOCX files...")
    print("=" * 50)
    
    python_info = analyze_docx(python_file)
    rust_info = analyze_docx(rust_file)
    
    print(f"Python version ({python_file}):")
    print(f"  Paragraphs: {python_info.get('paragraphs', 'ERROR')}")
    print(f"  Tables: {python_info.get('tables', 'ERROR')}")
    print(f"  Has header: {python_info.get('has_header', 'ERROR')}")
    print(f"  Content types: {sorted(python_info.get('content_types', []))}")
    
    print()
    print(f"Rust version ({rust_file}):")
    print(f"  Paragraphs: {rust_info.get('paragraphs', 'ERROR')}")
    print(f"  Tables: {rust_info.get('tables', 'ERROR')}")
    print(f"  Has header: {rust_info.get('has_header', 'ERROR')}")
    print(f"  Content types: {sorted(rust_info.get('content_types', []))}")
    
    print()
    print("Comparison:")
    
    # Check if key features match
    success = True
    
    if python_info.get('content_types') != rust_info.get('content_types'):
        print("❌ Content types don't match")
        print(f"   Python only: {python_info.get('content_types', set()) - rust_info.get('content_types', set())}")
        print(f"   Rust only: {rust_info.get('content_types', set()) - python_info.get('content_types', set())}")
        success = False
    else:
        print("✅ Content types match")
    
    # Check approximate paragraph count (allow some variance)
    python_paras = python_info.get('paragraphs', 0)
    rust_paras = rust_info.get('paragraphs', 0)
    para_diff = abs(python_paras - rust_paras)
    
    if para_diff <= 5:  # Allow small differences
        print(f"✅ Paragraph count similar ({python_paras} vs {rust_paras})")
    else:
        print(f"❌ Paragraph count differs significantly ({python_paras} vs {rust_paras})")
        success = False
    
    # Check table count
    if python_info.get('tables') == rust_info.get('tables'):
        print(f"✅ Table count matches ({python_info.get('tables')})")
    else:
        print(f"❌ Table count differs ({python_info.get('tables')} vs {rust_info.get('tables')})")
        success = False
    
    # Check header presence
    if python_info.get('has_header') == rust_info.get('has_header'):
        print("✅ Header presence matches")
    else:
        print("❌ Header presence differs")
        success = False
    
    if success:
        print("\n🎉 Overall: Files appear to have similar structure!")
    else:
        print("\n⚠️  Overall: Files have some differences that may need attention.")
    
    return 0 if success else 1

if __name__ == '__main__':
    sys.exit(main())